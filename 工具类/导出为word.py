import os
from docx import Document

def read_files_in_dir(directory, exclude_list=None):
    """ 递归读取目录下所有文件的内容，并排除指定的文件夹及其子文件夹中的所有文件 """
    if exclude_list is None:
        exclude_list = []

    result = []
    for root, dirs, files in os.walk(directory):
        # 排除指定的文件夹及其子文件夹
        dirs[:] = [d for d in dirs if not should_exclude(os.path.join(root, d), exclude_list)]

        level = root.replace(directory, '').count(os.sep)
        indent = ' ' * 4 * (level)
        result.append((f"{indent}{'-' * level} {os.path.basename(root)}/", ""))
        sub_indent = ' ' * 4 * (level + 1)
        for f in files:
            # 排除指定的文件夹及其子文件夹中的文件
            if should_exclude(os.path.join(root, f), exclude_list):
                continue

            filepath = os.path.join(root, f)
            try:
                with open(filepath, 'r', encoding='utf-8') as file:
                    content = file.read()
                    # 过滤掉非法字符
                    content = ''.join([c for c in content if 0x20 <= ord(c) <= 0xFF or ord(c) in (9, 10, 13)])
                    result.append((f"{sub_indent}{'-' * (level+1)} {f}:", content))
            except Exception as e:
                result.append((f"{sub_indent}{'-' * (level+1)} {f}: [无法读取文件]", ""))
    return result

def should_exclude(path, exclude_list):
    """ 判断路径是否应该被排除 """
    for exclude_path in exclude_list:
        if path.startswith(exclude_path):
            return True
    return False

def write_to_word(doc, items):
    """ 将文件名和内容写入Word文档 """
    for item in items:
        if item[0]:
            doc.add_paragraph(item[0])
        if item[1]:
            doc.add_paragraph(item[1])

if __name__ == "__main__":
    # directory = r'D:\1document\1test\PycharmProject_gitee\others\Projects\schedule'
    directory = r'D:\1document\1test\PycharmProject_gitee\others\Projects\rili'
    exclude_list = None
    # [
    #     r'D:\1document\1test\PycharmProject_gitee\others\Projects\schedule\node_modules'
    # ]

    # 读取文件夹中的所有文件，并排除指定的文件夹及其子文件夹中的所有文件
    items = read_files_in_dir(directory, exclude_list)

    # 创建Word文档
    doc = Document()

    # 写入内容
    write_to_word(doc, items)

    # 保存Word文档
    name = os.path.basename(directory)
    doc.save(name + '.docx')
    file_path = os.path.join(directory, name + '.docx')

    print(f"Word文档已创建,路径：{file_path}")
